"""Local file system connector for ingesting documents."""

import logging
import mimetypes
from datetime import datetime
from pathlib import Path
from typing import List, Optional

import markdown
from docx import Document as DocxDocument
from pypdf import PdfReader

from src.connectors.base import BaseConnector
from src.models import Document, DocumentMetadata, SourceType

logger = logging.getLogger(__name__)


class LocalFileConnector(BaseConnector):
    """Connector for reading documents from local file system."""

    # Supported file extensions
    SUPPORTED_EXTENSIONS = {
        ".txt",
        ".md",
        ".markdown",
        ".pdf",
        ".docx",
        ".doc",
    }

    def __init__(self) -> None:
        """Initialize local file connector."""
        super().__init__(SourceType.LOCAL)

    def validate_connection(self) -> bool:
        """Validate file system access.

        Returns:
            Always True for local file system
        """
        return True

    def fetch_documents(
        self, source_path: str, recursive: bool = True, **kwargs: any
    ) -> List[Document]:
        """Fetch documents from local directory or file.

        Args:
            source_path: Path to file or directory
            recursive: If True, recursively scan subdirectories
            **kwargs: Additional arguments (unused)

        Returns:
            List of documents
        """
        path = Path(source_path).expanduser().resolve()

        if not path.exists():
            logger.error(f"Path does not exist: {path}")
            return []

        documents = []

        if path.is_file():
            # Single file
            doc = self._load_file(path)
            if doc:
                documents.append(doc)
        elif path.is_dir():
            # Directory
            documents = self._load_directory(path, recursive)
        else:
            logger.warning(f"Path is neither file nor directory: {path}")

        logger.info(f"Loaded {len(documents)} documents from {source_path}")
        return documents

    def _load_directory(self, directory: Path, recursive: bool) -> List[Document]:
        """Load all supported documents from a directory.

        Args:
            directory: Directory path
            recursive: If True, scan subdirectories

        Returns:
            List of documents
        """
        documents = []
        pattern = "**/*" if recursive else "*"

        for file_path in directory.glob(pattern):
            if file_path.is_file() and file_path.suffix.lower() in self.SUPPORTED_EXTENSIONS:
                doc = self._load_file(file_path)
                if doc:
                    documents.append(doc)

        return documents

    def _load_file(self, file_path: Path) -> Optional[Document]:
        """Load a single file.

        Args:
            file_path: Path to file

        Returns:
            Document or None if loading failed
        """
        try:
            suffix = file_path.suffix.lower()

            if suffix in [".txt", ".md", ".markdown"]:
                content = self._read_text_file(file_path)
            elif suffix == ".pdf":
                content = self._read_pdf(file_path)
            elif suffix in [".docx", ".doc"]:
                content = self._read_docx(file_path)
            else:
                logger.warning(f"Unsupported file type: {suffix} for {file_path}")
                return None

            if not content or not content.strip():
                logger.warning(f"Empty content in file: {file_path}")
                return None

            # Get file metadata
            stat = file_path.stat()
            metadata = DocumentMetadata(
                source=str(file_path),
                source_type=self.source_type,
                title=file_path.stem,
                created_at=datetime.fromtimestamp(stat.st_ctime),
                modified_at=datetime.fromtimestamp(stat.st_mtime),
                file_type=suffix,
                file_size=stat.st_size,
            )

            return Document(content=content, metadata=metadata)

        except Exception as e:
            logger.error(f"Error loading file {file_path}: {e}")
            return None

    def _read_text_file(self, file_path: Path) -> str:
        """Read plain text or markdown file.

        Args:
            file_path: Path to file

        Returns:
            File content as string
        """
        with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
            content = f.read()

        # Convert markdown to plain text if needed
        if file_path.suffix.lower() in [".md", ".markdown"]:
            # Simple markdown to text conversion
            # More sophisticated conversion could use BeautifulSoup
            html = markdown.markdown(content)
            # Remove HTML tags (basic approach)
            import re

            content = re.sub(r"<[^>]+>", "", html)

        return content

    def _read_pdf(self, file_path: Path) -> str:
        """Read PDF file.

        Args:
            file_path: Path to PDF file

        Returns:
            Extracted text content
        """
        try:
            reader = PdfReader(str(file_path))
            text_parts = []

            for page in reader.pages:
                text = page.extract_text()
                if text:
                    text_parts.append(text)

            return "\n\n".join(text_parts)

        except Exception as e:
            logger.error(f"Error reading PDF {file_path}: {e}")
            return ""

    def _read_docx(self, file_path: Path) -> str:
        """Read DOCX file.

        Args:
            file_path: Path to DOCX file

        Returns:
            Extracted text content
        """
        try:
            doc = DocxDocument(str(file_path))
            text_parts = []

            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text)

            return "\n\n".join(text_parts)

        except Exception as e:
            logger.error(f"Error reading DOCX {file_path}: {e}")
            return ""
